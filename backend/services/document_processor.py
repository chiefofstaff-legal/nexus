"""
Document Processing Pipeline
=============================

Extract -> Classify -> File -> Index

Handles PDF (native + image-only), DOCX, and plain text documents.
Native text: PyMuPDF4LLM. Image-only: local Qwen2.5-VL via Ollama
(primary, on-prem) with Claude Haiku Vision as API fallback.
"""

import asyncio
import base64
import hashlib
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

from core.config import get_config
from core.idr_store import IDRStore
from core.ingestion_cost import record_call
from core.intent_decision_record import (
    DecisionPoint,
    IntentDecisionRecord,
    SynthesisMethod,
)
from models.document import (
    ClassificationResult,
    DocumentRecord,
    DocumentType,
    ExtractionResult,
    FilingResult,
)


class DocumentProcessor:
    """End-to-end document processing pipeline."""

    # Primary supported types with dedicated extractors
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".csv", ".json", ".xml", ".html", ".htm", ".rtf", ".eml"}

    # Text-like extensions that can be read directly
    TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".xml", ".html", ".htm", ".rtf", ".eml", ".log", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf", ".py", ".js", ".ts", ".java", ".c", ".cpp", ".h", ".go", ".rs", ".rb", ".sh", ".bat", ".sql"}

    def __init__(
        self,
        data_dir: Path,
        anthropic_client=None,
        idr_store: Optional[IDRStore] = None,
    ):
        self.data_dir = data_dir
        self.filed_dir = data_dir / "filed"
        self.filed_dir.mkdir(parents=True, exist_ok=True)
        self.anthropic_client = anthropic_client
        self.idr_store = idr_store

    async def extract_text(self, file_path: Path) -> ExtractionResult:
        """Extract text from a document file. Handles any readable file type."""
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return await self._extract_pdf(file_path)
        elif suffix == ".docx":
            return await self._extract_docx(file_path)
        elif suffix in self.TEXT_EXTENSIONS:
            text = file_path.read_text(encoding="utf-8", errors="replace")
            return ExtractionResult(
                text=text, page_count=1,
                method="plaintext", source_path=str(file_path),
            )
        else:
            # Try reading as text for any unknown extension
            try:
                text = file_path.read_text(encoding="utf-8", errors="replace")
                if text.strip():
                    return ExtractionResult(
                        text=text, page_count=1,
                        method="plaintext-fallback", source_path=str(file_path),
                    )
            except (UnicodeDecodeError, OSError):
                pass
            raise ValueError(f"Cannot extract text from: {suffix}")

    async def _extract_pdf(self, file_path: Path) -> ExtractionResult:
        """Extract text from PDF. Native text first, vision OCR on short output."""
        try:
            import pymupdf4llm
            text = pymupdf4llm.to_markdown(str(file_path))

            # If extracted text is too short, it's likely an image-only PDF
            if len(text.strip()) < 50:
                return await self._extract_pdf_via_vision(file_path)

            import pymupdf
            doc = pymupdf.open(str(file_path))
            page_count = len(doc)
            doc.close()

            return ExtractionResult(
                text=text, page_count=page_count,
                method="pymupdf4llm", source_path=str(file_path),
            )
        except Exception:
            return await self._extract_pdf_via_vision(file_path)

    def _render_pdf_pages_to_png(
        self, file_path: Path, max_pages: int, dpi: int = 150
    ) -> tuple[list[bytes], int]:
        """Render the first ``max_pages`` pages to PNG bytes. Returns (images, total_pages)."""
        import pymupdf
        doc = pymupdf.open(str(file_path))
        try:
            total_pages = len(doc)
            images: list[bytes] = []
            for page_num in range(min(max_pages, total_pages)):
                pix = doc[page_num].get_pixmap(dpi=dpi)
                images.append(pix.tobytes("png"))
            return images, total_pages
        finally:
            doc.close()

    async def _extract_pdf_via_vision(self, file_path: Path) -> ExtractionResult:
        """Vision OCR orchestrator: local Qwen primary -> Claude Vision fallback."""
        config = get_config()
        page_images, page_count = self._render_pdf_pages_to_png(
            file_path, max_pages=config.vision_ocr_max_pages
        )

        if not page_images:
            return ExtractionResult(
                text="[PDF has no pages to extract]",
                page_count=0, method="failed", source_path=str(file_path),
            )

        # Primary: local Qwen2.5-VL (on-prem, no PII leakage)
        if config.vision_ocr_enabled:
            try:
                text = await self._ocr_via_qwen_vl(
                    page_images,
                    model=config.vision_ocr_model,
                    base_url=config.ollama_base_url,
                    timeout_s=config.vision_ocr_timeout_s,
                )
                return ExtractionResult(
                    text=text, page_count=page_count,
                    method=f"qwen_vl:{config.vision_ocr_model}",
                    source_path=str(file_path),
                )
            except Exception:
                pass  # fall through to Claude

        # Fallback: Claude Haiku Vision (API, needs key)
        if not self.anthropic_client:
            return ExtractionResult(
                text="[PDF extraction requires local Qwen-VL or Claude API key]",
                page_count=page_count, method="failed", source_path=str(file_path),
            )

        text = await self._ocr_via_claude_vision(page_images)
        return ExtractionResult(
            text=text, page_count=page_count,
            method="claude_haiku_vision", source_path=str(file_path),
        )

    async def _ocr_via_qwen_vl(
        self, page_images: list[bytes], model: str, base_url: str, timeout_s: int
    ) -> str:
        """Extract text from pre-rendered page images via local Qwen2.5-VL."""
        import urllib.error
        import urllib.request

        prompt = (
            "Extract all text from this document page. "
            "Return only the extracted text, no commentary."
        )
        url = f"{base_url}/api/chat"

        def _call_one(img_bytes: bytes) -> str:
            img_b64 = base64.b64encode(img_bytes).decode()
            payload = json.dumps({
                "model": model,
                "messages": [
                    {"role": "user", "content": prompt, "images": [img_b64]},
                ],
                "stream": False,
                "options": {"temperature": 0.0},
            }).encode()
            req = urllib.request.Request(
                url, data=payload,
                headers={"Content-Type": "application/json"},
                method="POST",
            )
            with urllib.request.urlopen(req, timeout=timeout_s) as resp:
                data = json.loads(resp.read())
            return data.get("message", {}).get("content", "")

        texts: list[str] = []
        # Serial, not parallel -- M1 Pro 16 GB cannot afford two VL inferences
        # in flight. See H4 memory-pressure hypothesis in sprint plan.
        for img in page_images:
            page_text = await asyncio.to_thread(_call_one, img)
            texts.append(page_text)
        return "\n\n---\n\n".join(texts)

    async def _ocr_via_claude_vision(self, page_images: list[bytes]) -> str:
        """Extract text from pre-rendered page images via Claude Haiku Vision."""
        texts: list[str] = []
        for img_bytes in page_images:
            img_b64 = base64.b64encode(img_bytes).decode()
            response = await self.anthropic_client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=2000,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": img_b64}},
                        {"type": "text", "text": "Extract all text from this document page. Return only the extracted text, no commentary."},
                    ],
                }],
            )
            texts.append(response.content[0].text)
        return "\n\n---\n\n".join(texts)

    async def _extract_docx(self, file_path: Path) -> ExtractionResult:
        """Extract text from DOCX."""
        try:
            import docx
            doc = docx.Document(str(file_path))
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return ExtractionResult(
                text=text, page_count=len(doc.paragraphs) // 40 + 1,
                method="python-docx", source_path=str(file_path),
            )
        except ImportError:
            return ExtractionResult(
                text="[DOCX extraction requires python-docx]",
                page_count=0, method="failed", source_path=str(file_path),
            )

    async def classify(self, text: str, filename: str = "") -> ClassificationResult:
        """Classify a document using Claude Haiku."""
        if not self.anthropic_client:
            return self._classify_heuristic(text, filename)

        prompt = f"""Classify this legal document. Return ONLY valid JSON matching this schema:
{{
  "document_type": "contract|brief|correspondence|court_filing|invoice|memorandum|nda|dec|rse|rsg|other",
  "confidence": 0.0-1.0,
  "parties": ["Party A", "Party B"],
  "dates": ["2026-01-15"],
  "matter_reference": "case or matter number if found",
  "jurisdiction": "jurisdiction if identifiable",
  "summary": "one-sentence summary"
}}

Tier vocabulary (Swiss legal context — TODO: confirm semantics with Leandro):
- dec: Decision (Entscheidung / Décision / Decisione) — Swiss federal court ruling
- rse: Regulation (Recueil Systématique) — Federal collection of regulations
- rsg: Judgment (placeholder semantic) — to be confirmed with Sonnet Advisors

Filename: {filename}

Document text (first 3000 chars):
{text[:3000]}"""

        response = await self.anthropic_client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=500,
            messages=[{"role": "user", "content": prompt}],
        )

        # W3 cost meter — record Claude usage for ingestion classify calls.
        # Soft-fails: a logging error must not break ingestion.
        usage = getattr(response, "usage", None)
        if usage is not None:
            try:
                record_call(
                    model="claude-haiku-4-5-20251001",
                    input_tokens=getattr(usage, "input_tokens", 0) or 0,
                    output_tokens=getattr(usage, "output_tokens", 0) or 0,
                    purpose="classify",
                    document_id=filename,
                )
            except OSError:
                pass

        try:
            raw = response.content[0].text
            # Extract JSON from response
            json_match = re.search(r'\{[^{}]*\}', raw, re.DOTALL)
            if json_match:
                data = json.loads(json_match.group())
            else:
                data = json.loads(raw)

            return ClassificationResult(
                document_type=DocumentType(data.get("document_type", "other")),
                confidence=float(data.get("confidence", 0.5)),
                parties=data.get("parties", []),
                dates=data.get("dates", []),
                matter_reference=data.get("matter_reference"),
                jurisdiction=data.get("jurisdiction"),
                summary=data.get("summary", ""),
            )
        except (json.JSONDecodeError, ValueError):
            return self._classify_heuristic(text, filename)

    def _classify_heuristic(self, text: str, filename: str = "") -> ClassificationResult:
        """Fallback heuristic classification when no LLM available."""
        text_lower = text.lower()
        filename_lower = filename.lower()

        type_signals = {
            DocumentType.CONTRACT: ["agreement", "hereby", "terms and conditions", "party", "clause"],
            DocumentType.NDA: ["non-disclosure", "confidential", "nda", "proprietary information"],
            DocumentType.INVOICE: ["invoice", "amount due", "payment", "bill to", "total"],
            DocumentType.BRIEF: ["brief", "argument", "court", "plaintiff", "defendant"],
            DocumentType.COURT_FILING: ["filed", "court", "case no", "docket", "motion"],
            DocumentType.CORRESPONDENCE: ["dear", "sincerely", "regards", "re:"],
            DocumentType.MEMORANDUM: ["memorandum", "memo", "to:", "from:", "subject:"],
        }

        best_type = DocumentType.OTHER
        best_score = 0
        for doc_type, signals in type_signals.items():
            score = sum(1 for s in signals if s in text_lower or s in filename_lower)
            if score > best_score:
                best_score = score
                best_type = doc_type

        return ClassificationResult(
            document_type=best_type,
            confidence=min(0.3 + best_score * 0.15, 0.85),
            summary="Classified via heuristic (no LLM available)",
        )

    def generate_filing_path(
        self,
        original_path: Path,
        classification: ClassificationResult,
        user_id: str,
    ) -> FilingResult:
        """Generate auto-filed path scoped to ``user_id``.

        Layout: ``data/filed/<user_id>/<client>/<type>/<filename>``. The
        user_id prefix is what enforces tenant filesystem isolation —
        every other user's tree is reachable only by their own session.
        """
        if not user_id:
            raise ValueError("user_id is required for tenant-scoped filing")
        # Naming: {date}_{type}_{parties}_{hash}.{ext}
        date_str = datetime.utcnow().strftime("%Y-%m-%d")
        type_str = classification.document_type.value
        parties_str = "_".join(
            p.replace(" ", "-")[:20] for p in classification.parties[:2]
        ) or "unknown"
        content_hash = hashlib.sha256(original_path.name.encode()).hexdigest()[:8]
        ext = original_path.suffix

        new_filename = f"{date_str}_{type_str}_{parties_str}_{content_hash}{ext}"

        # Folder: /<user_id>/<client>/<type>/
        client = classification.parties[0].replace(" ", "-")[:30] if classification.parties else "unclassified"
        folder = self.filed_dir / user_id / client / type_str
        folder.mkdir(parents=True, exist_ok=True)

        new_path = folder / new_filename

        return FilingResult(
            original_path=str(original_path),
            new_filename=new_filename,
            new_path=str(new_path),
            document_type=classification.document_type,
            confidence=classification.confidence,
        )

    async def process(
        self,
        file_path: Path,
        user_id: str,
        extraction: Optional["ExtractionResult"] = None,
    ) -> DocumentRecord:
        """Full pipeline: extract -> classify -> file -> return record.

        Writes up to two IDRs if an ``idr_store`` was injected:

        - ``vision_ocr_provider``: which extractor actually served the
          file (pymupdf4llm / qwen_vl:<model> / claude_haiku_vision /
          plaintext / ...). Surfaces the on-prem vs cloud decision for
          FADP audits.
        - ``document_classification``: the per-document type label
          (contract | brief | nda | ...) with the filename as input
          summary. This satisfies H-MVP-3's "every decision writes an
          IDR" predicate for the ingestion surface.
        """
        if not user_id:
            raise ValueError("user_id is required for tenant-scoped processing")
        if extraction is None:
            extraction = await self.extract_text(file_path)
        classification = await self.classify(extraction.text, file_path.name)
        filing = self.generate_filing_path(file_path, classification, user_id=user_id)

        doc_id = hashlib.sha256(
            f"{file_path.name}:{datetime.utcnow().isoformat()}".encode()
        ).hexdigest()[:16]

        self._log_ingestion_idrs(file_path, extraction, classification, user_id)

        return DocumentRecord(
            id=doc_id,
            original_filename=file_path.name,
            filed_path=filing.new_path,
            document_type=classification.document_type,
            classification_confidence=classification.confidence,
            parties=classification.parties,
            dates=classification.dates,
            matter_reference=classification.matter_reference,
            jurisdiction=classification.jurisdiction,
            summary=classification.summary,
            page_count=extraction.page_count,
            extraction_method=extraction.method,
            chunk_count=0,  # Updated after indexing
        )

    def _log_ingestion_idrs(
        self,
        file_path: Path,
        extraction: ExtractionResult,
        classification: ClassificationResult,
        user_id: str,
    ) -> None:
        """Emit vision_ocr_provider + document_classification IDRs.

        ``user_id`` is the acting tenant (threaded from ``process``'s
        ``user_id`` arg, which the upload/batch/folder routes pass from
        ``current_user.id``). It is forwarded to ``IDRStore.append`` so
        the ingestion IDRs are tenant-attributed and visible on /idr.

        Silently no-ops when no IDR store was injected (keeps tests
        and lightweight consumers decoupled from the audit chain).
        Any IDR append failure is swallowed so ingestion never fails
        on audit infrastructure — the chain is a secondary concern
        relative to the primary file processing.
        """
        if self.idr_store is None:
            return
        input_hash = IntentDecisionRecord.hash_input(file_path.name)
        try:
            vision_idr = IntentDecisionRecord(
                decision_point=DecisionPoint.VISION_OCR_PROVIDER,
                input_hash=input_hash,
                input_summary=f"extraction provider for {file_path.name}",
                decision=extraction.method,
                confidence=1.0 if extraction.method != "failed" else 0.0,
                confidence_rationale=(
                    "deterministic: the extractor reports which path served"
                ),
                reasoning=(
                    f"{extraction.method} produced {extraction.page_count} "
                    f"pages and {len(extraction.text)} chars of text"
                ),
                synthesis_method=SynthesisMethod.DETERMINISTIC,
                falsification_criterion=(
                    "A second extraction of the same file with a different "
                    "provider (e.g. Qwen vs Claude vs Tesseract) would yield "
                    "materially different text content, indicating the "
                    "original method produced incomplete or hallucinated "
                    "output."
                ),
                metadata={
                    "extraction_method": extraction.method,
                    "page_count": extraction.page_count,
                    "text_len": len(extraction.text),
                    "filename": file_path.name,
                },
            )
            self.idr_store.append(vision_idr, user_id=user_id)
            doc_idr = IntentDecisionRecord(
                decision_point=DecisionPoint.DOCUMENT_CLASSIFICATION,
                input_hash=input_hash,
                input_summary=f"classification of {file_path.name}",
                decision=classification.document_type.value,
                confidence=classification.confidence,
                confidence_rationale=(
                    f"classifier returned {classification.confidence:.2f} on "
                    f"{classification.document_type.value}"
                ),
                reasoning=classification.summary
                or "no summary supplied by classifier",
                synthesis_method=SynthesisMethod.SINGLE_MODEL,
                falsification_criterion=(
                    "A legal-trained reviewer given the same file would "
                    "assign a different document_type. The specific parties "
                    f"and dates extracted ({', '.join(classification.parties[:3]) or 'none'} / "
                    f"{', '.join(classification.dates[:3]) or 'none'}) would "
                    "be judged inconsistent with the assigned type."
                ),
                metadata={
                    "filename": file_path.name,
                    "parties": classification.parties[:5],
                    "dates": classification.dates[:5],
                    "matter_reference": classification.matter_reference,
                    "jurisdiction": classification.jurisdiction,
                },
            )
            self.idr_store.append(doc_idr, user_id=user_id)
        except Exception:
            # Audit chain is secondary to ingestion success; don't fail
            # the whole pipeline because a signing key was misconfigured.
            pass
